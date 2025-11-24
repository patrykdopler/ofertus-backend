from fastapi import FastAPI, UploadFile, File
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from io import BytesIO
import zipfile


# -------------------------------------------------------
# FASTAPI + CORS (konieczne dla Netlify → Render)
# -------------------------------------------------------

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # można wpisać domenę Netlify
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# -------------------------------------------------------
# ODCZYT POZYCJI Z XLSX
# -------------------------------------------------------

def extract_positions(workbook):
    sheet = workbook.active
    max_row = sheet.max_row

    positions = []
    current = None

    for i in range(1, max_row + 1):
        c6 = sheet.cell(i, 6).value
        c7 = sheet.cell(i, 7).value

        if isinstance(c7, str) and c7.startswith("Poz."):
            if current:
                positions.append(current)

            current = {
                "lp": len(positions) + 1,
                "nazwa": c7,
                "ilosc": "",
                "opis": "",
                "image": None
            }

        if current is not None:
            if c6 == "Ilość:":
                current["ilosc"] = str(c7)
            elif c6 in ("Wypełnienia:", "Opis:"):
                if c7:
                    text = str(c7).replace("\n", " ").replace("_x000D_", " ")
                    current["opis"] += " " + text

    if current:
        positions.append(current)

    return positions


# -------------------------------------------------------
# WYCIĄGANIE OBRAZÓW Z XLSX (xl/media)
# -------------------------------------------------------

def extract_images(xlsx_bytes):
    images = []
    with zipfile.ZipFile(BytesIO(xlsx_bytes)) as z:
        for name in sorted(z.namelist()):
            if name.startswith("xl/media") and (
                name.lower().endswith(".jpg")
                or name.lower().endswith(".jpeg")
                or name.lower().endswith(".png")
            ):
                images.append(z.read(name))
    return images


# -------------------------------------------------------
# DODAWANIE OBRAZU DO KOMÓRKI POD NAZWĄ
# -------------------------------------------------------

def add_image_to_cell(cell, image_bytes, width_px=500):
    if not image_bytes:
        return

    width_inch = width_px / 96.0  # 96 dpi
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(BytesIO(image_bytes), width=Inches(width_inch))


# -------------------------------------------------------
# GŁÓWNY ENDPOINT: GENEROWANIE OFERTY
# -------------------------------------------------------

@app.post("/generate-offer")
async def generate_offer(xlsx: UploadFile = File(...)):
    xlsx_bytes = await xlsx.read()
    wb = load_workbook(BytesIO(xlsx_bytes), data_only=True)

    # Odczyt pozycji i zdjęć
    positions = extract_positions(wb)
    images = extract_images(xlsx_bytes)

    # Przypisanie zdjęć kolejno
    for i in range(len(positions)):
        if i < len(images):
            positions[i]["image"] = images[i]

    # Ładowanie template.docx
    doc = Document("template.docx")
    table = doc.tables[0]  # pierwsza tabela
    template_row = table.rows[1]

    # czyścimy starą zawartość
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)

    for pos in positions:
        row = table.add_row().cells
        row[0].text = str(pos["lp"])
        row[1].text = pos["nazwa"]
        if pos["image"]:
            add_image_to_cell(row[1], pos["image"])
        row[2].text = pos["ilosc"]
        row[3].text = pos["opis"]

    # generowanie DOCX
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=oferta.docx"}
    )


# -------------------------------------------------------
# ENDPOINT TESTOWY (opcjonalny)
# -------------------------------------------------------

@app.get("/")
def root():
    return {"status": "backend działa", "endpoint": "/generate-offer"}
