from fastapi import FastAPI, File, UploadFile
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from io import BytesIO
from datetime import datetime
import zipfile
import tempfile
import os

from docx import Document
from docx.shared import Cm
import openpyxl

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_items_from_xlsx(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[wb.sheetnames[0]]

    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append(list(row))

    items = []
    nazwa_cells = []
    wyp_cells = []

    for r_idx, row in enumerate(rows):
        if row is None:
            continue
        for c_idx, val in enumerate(row):
            text = str(val).strip() if val is not None else ""
            if text == "Nazwa:":
                nazwa_cells.append((r_idx, c_idx))
            if text in ("Wypełnienia:", "Wypełnienie:"):
                wyp_cells.append((r_idx, c_idx))

    def find_wypelnienie(after_row):
        for r, c in wyp_cells:
            if r > after_row:
                row = rows[r] or []
                if c + 1 < len(row):
                    val = row[c+1]
                    return str(val).strip() if val is not None else ""
        return ""

    for r, c in nazwa_cells:
        row = rows[r] or []
        nazwa = ""
        if c + 1 < len(row) and row[c+1] is not None:
            nazwa = str(row[c+1]).strip()
        if not nazwa:
            continue
        opis = find_wypelnienie(r)
        items.append({
            "lp": len(items) + 1,
            "nazwa": nazwa,
            "ilosc": 1,
            "opis": opis,
        })

    if not items:
        for row in rows:
            if not row:
                continue
            first = str(row[0]).strip() if row[0] is not None else ""
            if first.startswith("Poz"):
                opis = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
                items.append({
                    "lp": len(items) + 1,
                    "nazwa": first,
                    "ilosc": 1,
                    "opis": opis,
                })

    return items

def extract_images_from_xlsx(xlsx_path, tmp_dir):
    images_paths = []
    with zipfile.ZipFile(xlsx_path, "r") as zf:
        media_files = sorted([n for n in zf.namelist() if n.startswith("xl/media/")])
        for idx, name in enumerate(media_files):
            data = zf.read(name)
            ext = os.path.splitext(name)[1] or ".png"
            out_path = os.path.join(tmp_dir, f"image_{idx+1}{ext}")
            with open(out_path, "wb") as f:
                f.write(data)
            images_paths.append(out_path)
    return images_paths

def build_offer_doc(template_path, items, images_paths):
    doc = Document(template_path)

    today_str = datetime.now().strftime("%d.%m.%Y")
    for p in doc.paragraphs:
        if "{data}" in p.text:
            p.text = p.text.replace("{data}", today_str)

    table = None
    for t in doc.tables:
        if t.rows and t.rows[0].cells and "L.p." in (t.rows[0].cells[0].text or ""):
            table = t
            break

    if table is None:
        raise ValueError("Nie znaleziono tabeli z nagłówkiem 'L.p.' w szablonie DOCX")

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for i, item in enumerate(items):
        row = table.add_row()
        row.cells[0].text = str(item["lp"])
        row.cells[1].text = item["nazwa"]
        row.cells[2].text = str(item["ilosc"])
        row.cells[3].text = item["opis"] or ""

        if i < len(images_paths):
            img_row = table.add_row()
            merged = img_row.cells[0].merge(img_row.cells[1]).merge(img_row.cells[2]).merge(img_row.cells[3])
            paragraph = merged.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(images_paths[i], width=Cm(12))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

@app.post("/generate-offer")
async def generate_offer(file: UploadFile = File(...)):
    try:
        with tempfile.TemporaryDirectory() as tmp:
            xlsx_path = os.path.join(tmp, file.filename)
            with open(xlsx_path, "wb") as f:
                f.write(await file.read())

            items = extract_items_from_xlsx(xlsx_path)
            if not items:
                return JSONResponse(status_code=400, content={"detail": "Nie znaleziono pozycji w XLSX."})

            images_paths = extract_images_from_xlsx(xlsx_path, tmp)

            template_path = os.path.join(os.path.dirname(__file__), "Oferta_SZABLON4.docx")
            buf = build_offer_doc(template_path, items, images_paths)

            headers = {
                "Content-Disposition": 'attachment; filename="Oferta_Dopler.docx"'
            }
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers=headers,
            )
    except Exception as e:
        return JSONResponse(status_code=500, content={"detail": str(e)})
